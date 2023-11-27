import streamlit as st
from openai import OpenAI
from docx import Document
from io import BytesIO
from . import cine2013

# Initialize OpenAI client only once, not inside a function
client = OpenAI(api_key=st.secrets["openai"]["openai_api_key"])

def create_education_plan_doc(nombre_asignatura, campo_especifico, campo_detallado, topico, response_text):
    doc = Document()
    
    doc.add_heading('Planificación proceso de enseñanza y aprendizaje', level=1)

    intro_text = (f"A continuación, encontrará la planificación del proceso de enseñanza y aprendizaje diseñado "
                  f"por la Learnia para la asignatura {nombre_asignatura}, dentro del campo de conocimiento "
                  f"{campo_especifico}, {campo_detallado}, para los siguientes Resultados de Aprendizaje del contenido "
                  f"{topico}:\n")
    doc.add_paragraph(intro_text)

    doc.add_heading('Resultados de Aprendizaje, Indicadores de Logro y Metodologías de Aprendizaje Activo', level=2)
    doc.add_paragraph("A continuación, se presenta la propuesta de indicadores de logro y metodologías de aprendizaje "
                      "activo sugeridas para alcanzar los resultados de aprendizajes definidos:")
    doc.add_paragraph(response_text)

    doc.add_heading('Referencias', level=2)
    references = ("- Biggs J.B., & Collis K.F. (1982). Evaluating the Quality of Learning: The SOLO Taxonomy. New York: Academic Press.\n"
                  "- Campos de educación y capacitación 2013 de la CINE (ISCED-F 2013). Extraído desde https://uis.unesco.org/sites/default/files/documents/isced-fields-of-education-and-training-2013-sp.pdf\n"
                  "- Quality Matters, sitio web https://www.qualitymatters.org/")
    doc.add_paragraph(references)

    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Function to get chat response from GPT-4
def get_chat_response(nombre_asignatura, topico, campo_amplio, campo_especifico, campo_detallado):
    system_message = (
        f"Act as an expert analyst in the development and evaluation of educational curricula "
        f"focused on '{campo_amplio}', '{campo_especifico}', and '{campo_detallado}'."
    )
    user_request = (
        f"Considerando la asignatura '{nombre_asignatura}', propose three learning outcomes for "
        f"each level of the SOLO taxonomy for the topic '{topico}' using the structure Verb + Object + Context. "
        f"Ensure you provide 'Resultados de aprendizaje' for each of the 5 levels of the SOLO taxonomy, generate 3 achievement "
        f"indicators for each of the learning outcomes according to Quality Matters standards. Then, indicate the most "
        f"relevant active learning methodologies for collecting each of those indicators. Format the response as numbered multilevel lists. Remove ### and * symbols from the answers"
        f"All the answer should be in spanish language. "
    )

    prompt = f"{system_message}. For the request: {user_request}"

    try:
        response = client.chat.completions.create(
            model="gpt-4-1106-preview",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_request},
            ],
            max_tokens=4000,
            temperature=0.7,
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"An error occurred: {e}")
        return None

def display():
    # Using st.markdown with HTML and CSS to center the text
    st.markdown("""
        <style>
        .justified-text {
            text-align: center;
        }
        </style>
        <div class="centered-text">
        La app Learn-IA, basada en inteligencia artificial, optimiza el diseño instruccional al ayudar a crear resultados de aprendizaje, indicadores de logro y metodologías de aprendizaje activo. Los usuarios seleccionan el nombre de la asignatura, su contenido y campo de conocimiento según CINE 2013. La app proporciona propuestas de resultados de aprendizaje basadas en la taxonomía SOLO y sugiere indicadores de logro según los estándares de Quality Matters, además de metodologías de aprendizaje activo centradas en los estudiantes. Los usuarios pueden adaptar estas sugerencias para enriquecer el syllabus de su asignatura, adecuándolo al nivel deseado de conocimiento y profundidad. Espera por el archivo word para descargar.
        </div>
        """, unsafe_allow_html=True)
    
    # Initialize session state variables
    if 'response' not in st.session_state:
        st.session_state.response = None

    nombre_asignatura = st.text_input("Nombre de la Asignatura", key="nombre_asignatura_key")
    topico = st.text_input("Introduce un Contenido", key="topico_key")

    cmapos_amplios = list(cine2013.cine2013.keys())
    campo_amplio_seleccionado = st.selectbox("Selecciona un Campo Amplio", cmapos_amplios)
    campos_especificos = list(cine2013.cine2013[campo_amplio_seleccionado].keys())
    campo_especifico_seleccionado = st.selectbox("Selecciona un Campo Específico", campos_especificos)
    campos_detallados = cine2013.cine2013[campo_amplio_seleccionado][campo_especifico_seleccionado]
    campo_detallado_seleccionado = st.selectbox("Selecciona un Campo Detallado", campos_detallados)

    if st.button("Generar Resultados de Aprendizaje"):
        with st.spinner('Generando resultados de aprendizaje para el contenido seleccionado. Esto podría tomar unos minutos...'):
            response_text = get_chat_response(
                nombre_asignatura,
                topico,
                campo_amplio_seleccionado,
                campo_especifico_seleccionado,
                campo_detallado_seleccionado
            )

            if response_text:
                #st.write(response_text)

                doc_file = create_education_plan_doc(
                    nombre_asignatura=nombre_asignatura,
                    campo_especifico=campo_especifico_seleccionado,
                    campo_detallado=campo_detallado_seleccionado,
                    topico=topico,
                    response_text=response_text
                )
                
                # Sanitize the topic to create a valid file name
                #safe_topico = re.sub(r'[^\w\s]', '', topico).replace(' ', '_')

                st.download_button(
                    label="Descargar Planificación",
                    data=doc_file,
                    file_name="planificacion_curricular.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
